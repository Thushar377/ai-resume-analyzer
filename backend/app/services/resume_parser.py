# app/services/resume_parser.py
import PyPDF2
import docx
import io
import re
from typing import Dict, List, Any

class ResumeParser:
    
    # Common skills database
    TECH_SKILLS = {
        "programming": [
            "python", "javascript", "java", "c++", "c#", "ruby", "go", "rust",
            "typescript", "php", "swift", "kotlin", "scala", "r", "matlab"
        ],
        "web": [
            "react", "angular", "vue", "node.js", "django", "fastapi", "flask",
            "express", "spring boot", "next.js", "nuxt.js", "html", "css", "sass"
        ],
        "database": [
            "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "sqlite",
            "oracle", "cassandra", "dynamodb", "neo4j"
        ],
        "cloud": [
            "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible",
            "jenkins", "gitlab ci", "github actions", "heroku", "vercel"
        ],
        "data": [
            "machine learning", "deep learning", "tensorflow", "pytorch", "scikit-learn",
            "pandas", "numpy", "data analysis", "nlp", "computer vision", "tableau", "power bi"
        ],
        "soft_skills": [
            "leadership", "communication", "teamwork", "problem solving", "agile",
            "scrum", "project management", "critical thinking", "time management"
        ]
    }
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Extract from tables too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text."""
        text_lower = text.lower()
        found_skills = []
        
        for category, skills in self.TECH_SKILLS.items():
            for skill in skills:
                if skill.lower() in text_lower:
                    found_skills.append(skill.title())
        
        # Remove duplicates and return
        return list(set(found_skills))
    
    def extract_email(self, text: str) -> str:
        """Extract email from text."""
        pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(pattern, text)
        return matches[0] if matches else ""
    
    def extract_phone(self, text: str) -> str:
        """Extract phone number from text."""
        pattern = r'(\+?1?\s?)?(\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4})'
        matches = re.findall(pattern, text)
        return ''.join(matches[0]) if matches else ""
    
    def extract_linkedin(self, text: str) -> str:
        """Extract LinkedIn URL from text."""
        pattern = r'linkedin\.com/in/[\w-]+'
        matches = re.findall(pattern, text, re.IGNORECASE)
        return f"https://{matches[0]}" if matches else ""
    
    def extract_github(self, text: str) -> str:
        """Extract GitHub URL from text."""
        pattern = r'github\.com/[\w-]+'
        matches = re.findall(pattern, text, re.IGNORECASE)
        return f"https://{matches[0]}" if matches else ""
    
    def extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information."""
        return {
            "email": self.extract_email(text),
            "phone": self.extract_phone(text),
            "linkedin": self.extract_linkedin(text),
            "github": self.extract_github(text)
        }
    
    def extract_experience(self, text: str) -> Dict[str, Any]:
        """Extract work experience information."""
        years_pattern = r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp)'
        years_matches = re.findall(years_pattern, text, re.IGNORECASE)
        
        # Common job titles
        job_titles = [
            "Software Engineer", "Developer", "Manager", "Analyst",
            "Designer", "Architect", "Lead", "Director", "CTO", "CEO",
            "Data Scientist", "DevOps", "Full Stack", "Backend", "Frontend"
        ]
        
        found_titles = [title for title in job_titles if title.lower() in text.lower()]
        total_years = max([int(y) for y in years_matches], default=0)
        
        return {
            "total_years": total_years,
            "positions": found_titles[:5],
            "has_management": any(
                kw in text.lower() for kw in ["managed", "led", "supervised", "directed"]
            )
        }
    
    def extract_education(self, text: str) -> Dict[str, Any]:
        """Extract education information."""
        degrees = {
            "phd": ["phd", "ph.d", "doctorate"],
            "masters": ["master", "m.s.", "m.sc", "mba", "m.eng"],
            "bachelors": ["bachelor", "b.s.", "b.sc", "b.e.", "b.tech", "b.eng"],
            "associate": ["associate", "a.s.", "a.a."]
        }
        
        universities_keywords = ["university", "college", "institute", "school of"]
        
        found_degree = None
        for degree, keywords in degrees.items():
            if any(kw in text.lower() for kw in keywords):
                found_degree = degree
                break
        
        has_university = any(kw in text.lower() for kw in universities_keywords)
        
        # Extract GPA
        gpa_pattern = r'GPA[:\s]+(\d+\.\d+)'
        gpa_matches = re.findall(gpa_pattern, text, re.IGNORECASE)
        
        return {
            "highest_degree": found_degree,
            "has_university": has_university,
            "gpa": gpa_matches[0] if gpa_matches else None
        }
    
    def parse_resume(self, file_content: bytes, file_type: str) -> Dict[str, Any]:
        """Main method to parse resume and extract all information."""
        if file_type == "pdf":
            raw_text = self.extract_text_from_pdf(file_content)
        elif file_type == "docx":
            raw_text = self.extract_text_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return {
            "raw_text": raw_text,
            "skills": self.extract_skills(raw_text),
            "contact": self.extract_contact_info(raw_text),
            "experience": self.extract_experience(raw_text),
            "education": self.extract_education(raw_text)
        }

resume_parser = ResumeParser()